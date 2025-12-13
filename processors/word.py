"""Word document processing utilities."""

import os
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


def process(file_path, use_ocr=None):
    """
    Extract text from Word document and convert to Markdown.

    Args:
        file_path: Path to .docx file
        use_ocr: Ignored for Word docs (always direct extraction)

    Returns:
        {
            "text": str,
            "file_size": int,
            "num_pages": int,
            "processing_method": str,
            "metadata": dict
        }
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word file not found: {file_path}")

    file_size = os.path.getsize(file_path)

    try:
        doc = Document(file_path)
    except Exception as e:
        raise ValueError(f"Invalid Word file: {e}")

    md_parts = []
    paragraph_count = 0
    table_count = 0

    for element in iter_block_items(doc):
        if isinstance(element, Paragraph):
            para = element
            text = para.text.strip()

            if not text:
                continue

            paragraph_count += 1

            style_name = para.style.name if para.style else ''

            if 'Heading 1' in style_name or 'Title' in style_name:
                md_parts.append(f"# {text}")
            elif 'Heading 2' in style_name:
                md_parts.append(f"## {text}")
            elif 'Heading 3' in style_name:
                md_parts.append(f"### {text}")
            elif 'Heading 4' in style_name:
                md_parts.append(f"#### {text}")
            elif 'List' in style_name or 'Bullet' in style_name:
                md_parts.append(f"- {text}")
            else:
                md_parts.append(text)

        elif isinstance(element, Table):
            table = element
            table_count += 1
            md_table = convert_table_to_markdown(table)
            if md_table:
                md_parts.append(md_table)

    full_text = '\n\n'.join(md_parts)

    if not full_text.strip():
        raise ValueError("Could not extract any text from Word document.")

    return {
        "text": full_text,
        "file_size": file_size,
        "num_pages": estimate_page_count(len(full_text)),
        "processing_method": "direct",
        "metadata": {
            "paragraphs": paragraph_count,
            "tables": table_count
        }
    }


def iter_block_items(document):
    """Iterate through paragraphs and tables in document order."""
    from docx.document import Document as DocxDocument
    from docx.oxml.ns import qn

    body = document.element.body

    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, document)
        elif child.tag == qn('w:tbl'):
            yield Table(child, document)


def convert_table_to_markdown(table):
    """Convert Word table to Markdown table."""
    rows = []

    for row in table.rows:
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    max_cols = max(len(row) for row in rows)
    for row in rows:
        while len(row) < max_cols:
            row.append("")

    md_lines = []

    header = rows[0]
    md_lines.append('| ' + ' | '.join(header) + ' |')
    md_lines.append('|' + '|'.join(['---'] * max_cols) + '|')

    for row in rows[1:]:
        md_lines.append('| ' + ' | '.join(row) + ' |')

    return '\n'.join(md_lines)


def estimate_page_count(char_count):
    """Estimate page count from character count (roughly 3000 chars/page)."""
    return max(1, char_count // 3000)


def analyze_word(file_path):
    """
    Analyze Word document.

    Args:
        file_path: Path to Word file

    Returns:
        {
            "num_pages": int,
            "ocr_recommended": bool,
            "ocr_reason": str
        }
    """
    try:
        doc = Document(file_path)
        total_chars = sum(len(p.text) for p in doc.paragraphs)

        return {
            "num_pages": estimate_page_count(total_chars),
            "ocr_recommended": False,
            "ocr_reason": "Word documents use direct text extraction"
        }
    except Exception as e:
        return {
            "num_pages": 0,
            "ocr_recommended": False,
            "ocr_reason": f"Analysis failed: {str(e)}"
        }
